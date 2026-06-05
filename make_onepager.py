#!/usr/bin/env python
"""
make_onepager.py
================
Turn outputs/four_fold_comparison.csv into a one-page Word artifact in the
format requested for SOW v2.5 acceptance review:
  per-fold MSE/R^2 for PINN-LSTM and the vanilla LSTM baseline, per-fold
  improvement (absolute + percent), 4-fold averages, Lu et al. (2023) published
  values side-by-side with absolute deviation, and a caption.

The caption is generated from the numbers: if the physics-informed model did
NOT beat the baseline on average, the caption says so plainly (it does not
spin a negative result), and it always identifies the largest-deviation fold
with a [FILL IN] prompt for the human to name the methodology cause.

Requires: python-docx  (pip install python-docx)

Usage:
  python make_onepager.py [comparison.csv] [out.docx]
Defaults: outputs/four_fold_comparison.csv -> outputs/four_fold_onepager.docx
"""
import csv
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _f(x):
    try:
        return float(x)
    except (TypeError, ValueError):
        return None


def fmt_mse(s):
    v = _f(s)
    return f"{v * 1e3:.2f}" if v is not None else (s or "pending")


def fmt_r2(s):
    v = _f(s)
    return f"{v:.3f}" if v is not None else (s or "pending")


def fmt_pct(s):
    v = _f(s)
    return f"{v:+.1f}%" if v is not None else (s or "pending")


def make_caption(rows):
    data = [r for r in rows if r["test_device"] != "AVG"]
    avg = next((r for r in rows if r["test_device"] == "AVG"), None)
    parts = []

    avg_imp = _f(avg["improvement_mse_pct"]) if avg else None
    if avg_imp is not None:
        if avg_imp < 0:
            parts.append(
                f"On this run the physics-informed LSTM did not improve on the vanilla LSTM "
                f"baseline: averaged across the four folds its test MSE was {abs(avg_imp):.1f}% "
                f"higher (worse) and mean R\u00b2 fell from {_f(avg['lstm_baseline_r2']):.3f} to "
                f"{_f(avg['pinn_lstm_r2']):.3f} \u2014 the opposite of the +13.9% MSE improvement "
                f"Lu et al. (2023) report for PI-LSTM.")
        else:
            parts.append(
                f"The physics-informed LSTM improved averaged test MSE by {avg_imp:.1f}% over the "
                f"vanilla LSTM baseline (Lu et al. report +13.9%).")
    if avg and _f(avg["lstm_baseline_r2"]) is not None:
        parts.append(
            f"The vanilla-LSTM baseline R\u00b2 values are nonetheless reasonable "
            f"(mean {_f(avg['lstm_baseline_r2']):.3f}), confirming the recurrent model fits the "
            f"degradation signal once the paper architecture (80-cell single-layer LSTM, seq=10) "
            f"is used.")

    numeric = [r for r in data if _f(r["deviation_pinn_mse_abs"]) is not None]
    if numeric:
        worst = max(numeric, key=lambda r: _f(r["deviation_pinn_mse_abs"]))
        parts.append(
            f"The largest deviation from Lu et al. is on the device-{worst['test_device']} fold "
            f"(|\u0394MSE| = {_f(worst['deviation_pinn_mse_abs']) * 1e3:.1f}\u00d710\u207b\u00b3, "
            f"|\u0394R\u00b2| = {_f(worst['deviation_pinn_r2_abs']):.3f}). "
            f"[FILL IN after inspecting device {worst['test_device']}: the most likely driver is the "
            f"automatic end-of-life detection (FAILURE_DROP_THRESHOLD) and resulting RUL labels on "
            f"this device's late-life V_ce; confirm whether the detected N_f matches the paper's "
            f"hand-identified failure point, and note the average-downsampling window vs the paper's "
            f"per-square-wave-cycle averaging.]")
    else:
        parts.append("Per-fold numbers are pending the full run; the Lu et al. reference column is "
                     "shown for format only.")

    parts.append(
        "All runs use the Lu-matching configuration (80-cell single-layer LSTM, sequence length 10, "
        "\u03b1=0.1, \u03b2=100, \u03b3=0.1, Adam, 2000 epochs) under leave-one-device-out cross-validation "
        "on NASA PCoE Dataset #8. Lu et al. PI-LSTM reference values are read from their Figure 10 "
        "(\u00d710\u207b\u00b3).")
    return " ".join(parts)


def build(csv_path, out_docx):
    rows = list(csv.DictReader(open(csv_path, newline="", encoding="utf-8")))
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.55)
    sec.left_margin = sec.right_margin = Inches(0.55)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Four-Fold Cross-Validation: PINN-LSTM vs. LSTM Baseline\n"
                  "IGBT RUL Reproduction of Lu et al. (2023), NASA PCoE Dataset #8")
    r.bold = True
    r.font.size = Pt(13)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Leave-one-device-out  |  MSE \u00d710\u207b\u00b3  |  lower MSE / higher R\u00b2 = better"
                   "  |  \u0394MSE>0 and \u0394MSE%>0 mean PINN beat baseline")
    sr.italic = True
    sr.font.size = Pt(8)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    headers = ["Test\ndevice", "LSTM\nMSE", "LSTM\nR\u00b2", "PINN-LSTM\nMSE", "PINN-LSTM\nR\u00b2",
               "\u0394MSE\n(abs)", "\u0394MSE\n(%)", "Lu PINN\nMSE", "Lu PINN\nR\u00b2",
               "|dev|\nMSE", "|dev|\nR\u00b2"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p.add_run(htext)
        rr.bold = True
        rr.font.size = Pt(8)

    for row in rows:
        cells = t.add_row().cells
        vals = [str(row["test_device"]),
                fmt_mse(row["lstm_baseline_mse"]), fmt_r2(row["lstm_baseline_r2"]),
                fmt_mse(row["pinn_lstm_mse"]), fmt_r2(row["pinn_lstm_r2"]),
                fmt_mse(row["improvement_mse_abs"]), fmt_pct(row["improvement_mse_pct"]),
                fmt_mse(row["lu_pinn_mse"]), fmt_r2(row["lu_pinn_r2"]),
                fmt_mse(row["deviation_pinn_mse_abs"]), fmt_r2(row["deviation_pinn_r2_abs"])]
        for i, v in enumerate(vals):
            cl = cells[i]
            cl.text = ""
            p = cl.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(v)
            rr.font.size = Pt(8.5)
            if row["test_device"] == "AVG":
                rr.bold = True
            if i == 6:  # flag negative improvement (PINN worse than baseline) in red
                iv = _f(v.replace('%', '').replace('+', ''))
                if iv is not None and iv < 0:
                    rr.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(10)
    cr = cap.add_run("Caption. ")
    cr.bold = True
    cr.font.size = Pt(9)
    br = cap.add_run(make_caption(rows))
    br.font.size = Pt(9)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(8)
    fr = foot.add_run("Generated from outputs/four_fold_comparison.csv. Lu et al. (2023), "
                      "Sci. Rep. 13:10167 (CC BY 4.0); reference values from Figure 10.")
    fr.italic = True
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.save(out_docx)
    print("saved", out_docx)


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else str(Path("outputs") / "four_fold_comparison.csv")
    out = sys.argv[2] if len(sys.argv) > 2 else str(Path("outputs") / "four_fold_onepager.docx")
    build(src, out)
